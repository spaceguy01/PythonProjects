""" Working with EXISTING WORD documents using Python docx """

import docx

""" Open Existing WORD document and set it to a variable """
doc = docx.Document('Word.docx')

""" Get Individual Paragraph in WORD doc or Loop to get all Paragraph Text """
doc.paragraphs[0].text
doc.paragraphs[1].text

for each in doc.paragraphs:
    print (each.text)

""" Set a paragraph to a variable and check .runs for each paragraph """
""" paragraph.runs shows what styles are in the paragraph """
p1 = doc.paragraphs[0]
p1.runs[0]

""" To get the text for each run in paragraph, use .text """
p1.runs[0].text

for each in p1.runs:
    print (each.text)

""" Boolean to check if something is bold/italic/underline for example """
p1.runs[1].bold
p1.runs[1].underline
p1.runs[1].italic

""" Changing the Style of the runs """
p1.runs[1].bold == True    """ Sets runs[1] as bold """
p1.runs[1].underline == False """ Removes the underline from runs[1]"""
p1.runs[1].italic == None """ Sets italic as None for runs[1]"""

""" Changing the style of the paragraph """
p1.style = 'Title'   """ Makes p1 paragraph into Title style """


"""-------------------------------------------------------"""

""" Creating New WORD document """
doc = docx.Document()

""" Add a paragraph to the New WORD document """
doc.add_paragraph('text to input')

""" Save the modified WORD file """
doc.save('newfilename.docx')
